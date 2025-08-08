from flask import Flask, render_template, request, send_file, redirect, url_for, flash, abort
from docx import Document
from docx.shared import RGBColor
from docx.table import Table
import os, re, io
from datetime import datetime
from zipfile import ZipFile
import pandas as pd

# ====== НАСТРОЙКИ ======
TEMPLATE_DIR = "templates_docs"        # где лежат .docx-шаблоны
GENERATED_DIR = "generated"            # куда сохраняем результаты
ALLOWED_SUFFIX = ".docx"

# Excel с артистами/треками
CATALOG_PATH = "data/report.xlsx"
# Шаблон приложения 1 (точное имя файла в templates_docs/)
APPENDIX_TEMPLATE_NAME = "приложение 1 — копия.docx"

os.makedirs(GENERATED_DIR, exist_ok=True)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret")

# ====== ДОГОВОРА: бережная замена плейсхолдеров ======
PLACEHOLDER_RE = re.compile(r"\{[^{}]+\}")

def list_templates():
    return sorted([f for f in os.listdir(TEMPLATE_DIR) if f.lower().endswith(ALLOWED_SUFFIX)])

def extract_placeholders(docx_paths):
    """Собираем уникальные плейсхолдеры из тела, таблиц и колонтитулов."""
    found = set()
    for path in docx_paths:
        doc = Document(path)
        # основной текст
        for p in doc.paragraphs:
            full = "".join(run.text for run in p.runs) or p.text
            if full:
                found.update(PLACEHOLDER_RE.findall(full))
        # таблицы
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full = "".join(run.text for run in p.runs) or p.text
                        if full:
                            found.update(PLACEHOLDER_RE.findall(full))
        # колонтитулы
        for section in doc.sections:
            for p in section.header.paragraphs:
                full = "".join(run.text for run in p.runs) or p.text
                if full:
                    found.update(PLACEHOLDER_RE.findall(full))
            for p in section.footer.paragraphs:
                full = "".join(run.text for run in p.runs) or p.text
                if full:
                    found.update(PLACEHOLDER_RE.findall(full))
    return sorted(found)

def replace_placeholders_preserve_runs(paragraph, mapping):
    """
    Бережная замена: сохраняем исходные runs (шрифты/размеры/стили).
    Меняем ТОЛЬКО текст плейсхолдера на value и красим его в чёрный.
    """
    if not paragraph.runs:
        return
    run_texts = [r.text or "" for r in paragraph.runs]
    full = "".join(run_texts)
    if not full:
        return
    matches = list(PLACEHOLDER_RE.finditer(full))
    if not matches:
        return
    lengths = [len(t) for t in run_texts]
    cumul, s = [], 0
    for L in lengths:
        cumul.append(s); s += L
    def locate(pos: int):
        i = 0
        while i + 1 < len(cumul) and cumul[i + 1] <= pos:
            i += 1
        return i, pos - cumul[i]
    for m in reversed(matches):
        ph_text = m.group(0)
        if ph_text not in mapping:
            continue
        value = mapping[ph_text]
        si, so = locate(m.start())
        ei, eo = locate(m.end() - 1)
        if si == ei:
            r = paragraph.runs[si]
            t = r.text or ""
            r.text = t[:so] + value + t[eo + 1:]
            try: r.font.color.rgb = RGBColor(0, 0, 0)
            except: pass
        else:
            r_start, r_end = paragraph.runs[si], paragraph.runs[ei]
            t_start, t_end = (r_start.text or ""), (r_end.text or "")
            r_start.text = t_start[:so] + value
            try: r_start.font.color.rgb = RGBColor(0, 0, 0)
            except: pass
            r_end.text = t_end[eo + 1:]
            for idx in range(si + 1, ei):
                paragraph.runs[idx].text = ""

def replace_in_doc(doc, mapping):
    for p in doc.paragraphs:
        replace_placeholders_preserve_runs(p, mapping)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_preserve_runs(p, mapping)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_placeholders_preserve_runs(p, mapping)
        for p in section.footer.paragraphs:
            replace_placeholders_preserve_runs(p, mapping)

# ====== EXCEL → ПРИЛОЖЕНИЕ 1 ======
_catalog_df = None
def load_catalog():
    """Читаем Excel и нормализуем строки."""
    global _catalog_df
    if _catalog_df is None:
        if not os.path.exists(CATALOG_PATH):
            raise FileNotFoundError(f"Excel не найден: {CATALOG_PATH}")
        df = pd.read_excel(CATALOG_PATH)
        for c in df.columns:
            if df[c].dtype == object:
                df[c] = df[c].astype(str).fillna("").str.strip()
        _catalog_df = df
    return _catalog_df

def rows_for_artist(artist_name: str):
    """Находим все треки артиста (учитывает 'фиты', если они в artist_name)."""
    df = load_catalog()
    q = re.sub(r"\s+", " ", (artist_name or "")).strip().lower()
    if not q:
        return df.iloc[0:0]
    mask = df["artist_name"].str.lower().str.contains(q, na=False)
    rows = df[mask].copy()
    sort_cols = [c for c in ["album_name", "track_name"] if c in rows.columns]
    if sort_cols:
        rows = rows.sort_values(by=sort_cols)
    return rows

def get_appendix_table(doc: Document) -> Table:
    """Берём первую таблицу в шаблоне приложения (если другая — скажи, сделаем поиск по маркеру)."""
    if not doc.tables:
        raise RuntimeError("В шаблоне приложения нет таблиц.")
    return doc.tables[0]

def clear_cell(cell):
    for p in cell.paragraphs:
        p.text = ""

def set_row_cells(row, values):
    for i, val in enumerate(values):
        if i < len(row.cells):
            clear_cell(row.cells[i])
            row.cells[i].paragraphs[0].add_run("" if val is None else str(val))

def clone_row(table: Table, template_idx: int):
    """Клонируем строку таблицы (переносим стиль первого параграфа)."""
    new = table.add_row()
    for i, cell in enumerate(table.rows[template_row_idx].cells):
        if i < len(new.cells):
            if cell.paragraphs and new.cells[i].paragraphs:
                new.cells[i].paragraphs[0].style = cell.paragraphs[0].style
    return new

def map_record_to_values(idx: int, rec: dict):
    """
    Порядок колонок "Приложения 1":
    1) № п/п
    2) Исполнитель            -> artist_name
    3) Композиция             -> track_name
    4) Автор слов             -> lyricist
    5) Автор музыки           -> composer
    6) Альбом                 -> album_name
    7) Доля авторских прав    -> author_right
    8) Доля смежных прав      -> related_right
    9) Территория             -> countries (если пусто — "Весь мир")
    10) Копирайт              -> label (если пусто — "Divine Music")
    11) Примечание            -> rec["note"] если появится, иначе пусто
    """
    artist      = rec.get("artist_name", "")
    track       = rec.get("track_name", "")
    lyrics      = rec.get("lyricist", "")
    music       = rec.get("composer", "")
    album       = rec.get("album_name", "")
    auth_share  = rec.get("author_right", "")
    neigh_share = rec.get("related_right", "")
    territory   = rec.get("countries", "") or "Весь мир"
    copyright_  = rec.get("label", "") or "Divine Music"
    note        = rec.get("note", "")
    return [str(idx), artist, track, lyrics, music, album, auth_share, neigh_share, territory, copyright_, note]

def fill_appendix_table(doc: Document, df: pd.DataFrame):
    """Заполняем таблицу Приложения 1, нумеруем строки с 1, расширяем на любое кол-во треков."""
    table = get_appendix_table(doc)
    # Часто: 1-я строка — заголовки, 2-я — строка-образец
    global template_row_idx
    template_row_idx = 1 if len(table.rows) > 1 else 0

    rows = df.to_dict(orient="records")
    if not rows:
        return

    # первая запись в шаблонной строке с №=1
    set_row_cells(table.rows[template_row_idx], map_record_to_values(1, rows[0]))

    # остальные — клонируем образец (№ = 2..N)
    for i, rec in enumerate(rows[1:], start=2):
        r = clone_row(table, template_row_idx)
        set_row_cells(r, map_record_to_values(i, rec))

# ====== РОУТЫ: главная (договоры) ======
@app.route("/")
def index():
    templates = list_templates()
    selected = request.args.getlist("t")
    selected = [s for s in selected if s in templates]
    placeholders = extract_placeholders([os.path.join(TEMPLATE_DIR, f) for f in selected]) if selected else []
    return render_template("index.html", templates=templates, selected=selected, placeholders=placeholders)

@app.route("/placeholders")
def placeholders():
    templates = list_templates()
    selected = request.args.getlist("t")
    selected = [s for s in selected if s in templates]
    placeholders = extract_placeholders([os.path.join(TEMPLATE_DIR, f) for f in selected]) if selected else []
    return render_template("index.html", templates=templates, selected=selected, placeholders=placeholders)

@app.route("/generate", methods=["POST"])
def generate():
    templates = list_templates()
    selected = request.form.getlist("selected_templates")
    selected = [s for s in selected if s in templates]
    if not selected:
        flash("Выберите хотя бы один шаблон.", "warning")
        return redirect(url_for("index"))

    mapping = {k[3:]: v for k, v in request.form.items() if k.startswith("ph:")}
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_files = []

    # Имя по {ФИО}
    fio_value = (mapping.get("{ФИО}", "") or "").strip() or "БезФИО"
    safe_fio = re.sub(r'[\\/*?:"<>|]', "", fio_value)

    for name in selected:
        path = os.path.join(TEMPLATE_DIR, name)
        doc = Document(path)
        replace_in_doc(doc, mapping)

        base = os.path.splitext(name)[0]
        base_lower = base.lower()
        if "договор" in base_lower:
            kind = "договор"
        elif "приложение" in base_lower:
            kind = "приложение"
        else:
            parts = re.split(r"[\s_()-]+", base)
            kind = (parts[-1] if parts and parts[-1] else base)

        out_name = f"{safe_fio} {kind}.docx"
        out_path = os.path.join(GENERATED_DIR, out_name)
        doc.save(out_path)
        out_files.append(out_path)

    mem = io.BytesIO()
    with ZipFile(mem, "w") as z:
        for fp in out_files:
            z.write(fp, arcname=os.path.basename(fp))
    mem.seek(0)
    zip_name = f"{safe_fio}.zip" if safe_fio else f"contracts_{stamp}.zip"
    return send_file(mem, as_attachment=True, download_name=zip_name)

# ====== РОУТЫ: Приложение 1 ======
@app.route("/appendix", methods=["GET"])
def appendix():
    artist_q = request.args.get("artist", "").strip()
    tracks = None
    placeholders = []
    # плейсхолдеры из шаблона приложения
    appx_path = os.path.join(TEMPLATE_DIR, APPENDIX_TEMPLATE_NAME)
    if os.path.exists(appx_path):
        placeholders = extract_placeholders([appx_path])

    if artist_q:
        try:
            tracks = rows_for_artist(artist_q).reset_index(drop=True)  # ВАЖНО
        except Exception as e:
            flash(f"Ошибка чтения Excel: {e}", "warning")
    return render_template("appendix.html", artist=artist_q, tracks=tracks, placeholders=placeholders)

@app.route("/appendix/generate", methods=["POST"])
def appendix_generate():
    artist_q = request.form.get("artist", "").strip()
    sel_idx = request.form.getlist("sel")

    rows = rows_for_artist(artist_q).reset_index(drop=True)  # ВАЖНО

    # безопасно парсим индексы из чекбоксов
    try:
        idxs = [int(i) for i in sel_idx]
        idxs = [i for i in idxs if 0 <= i < len(rows)]
    except Exception:
        idxs = []
    if idxs:
        rows = rows.iloc[idxs]

    appx_path = os.path.join(TEMPLATE_DIR, APPENDIX_TEMPLATE_NAME)
    if not os.path.exists(appx_path):
        flash(f"Не найден шаблон приложения: {APPENDIX_TEMPLATE_NAME}", "warning")
        return redirect(url_for("appendix", artist=artist_q))

    # собираем значения для плейсхолдеров из формы (поля вида name="ph:{...}")
    mapping = {k[3:]: v for k, v in request.form.items() if k.startswith("ph:")}

    doc = Document(appx_path)
    # сначала подставим плейсхолдеры в тексте
    replace_in_doc(doc, mapping)
    # затем заполним таблицу треками
    fill_appendix_table(doc, rows)

    # имя файла по ФИО (если передали)
    fio_value = (mapping.get("{ФИО}", "") or "").strip() or "БезФИО"
    safe_fio = re.sub(r'[\\/*?:"<>|]', "", fio_value)
    out_name = f"{safe_fio} приложение.docx"
    out_path = os.path.join(GENERATED_DIR, out_name)
    doc.save(out_path)

    return send_file(out_path, as_attachment=True, download_name=out_name)

# ====== ПРОЧЕЕ ======
@app.route("/downloads")
def downloads():
    files = sorted([f for f in os.listdir(GENERATED_DIR) if f.lower().endswith(".docx")], reverse=True)
    return render_template("downloads.html", files=files)

@app.route("/download/<name>")
def download_file(name):
    if not name.lower().endswith(".docx"):
        abort(404)
    path = os.path.join(GENERATED_DIR, name)
    if not os.path.exists(path):
        abort(404)
    return send_file(path, as_attachment=True, download_name=name)

@app.errorhandler(404)
def not_found(e):
    return render_template("error.html", code=404, message="Страница не найдена"), 404

@app.errorhandler(500)
def server_error(e):
    return render_template("error.html", code=500, message="Внутренняя ошибка сервера"), 500
    # ====== Мастер: договор(ы)+приложение на одной странице ======
@app.route("/all-in-one", methods=["GET"])
def all_in_one():
    templates = list_templates()  # список .docx в templates_docs
    selected = request.args.getlist("t")
    selected = [s for s in selected if s in templates]

    # флаг: включено ли приложение
    appendix_on = request.args.get("appendix") == "on"
    artist_q = request.args.get("artist", "").strip()

    # соберём плейсхолдеры: из выбранных договоров + (опц.) из шаблона приложения
    ph_paths = [os.path.join(TEMPLATE_DIR, f) for f in selected]
    if appendix_on:
        appx_path = os.path.join(TEMPLATE_DIR, APPENDIX_TEMPLATE_NAME)
        if os.path.exists(appx_path):
            ph_paths.append(appx_path)
    placeholders = extract_placeholders(ph_paths) if ph_paths else []

    # если включено приложение и введён артист — подтянем треки
    tracks = None
    if appendix_on and artist_q:
        try:
            tracks = rows_for_artist(artist_q).reset_index(drop=True)
        except Exception as e:
            flash(f"Ошибка чтения Excel: {e}", "warning")

    return render_template(
        "all_in_one.html",
        templates=templates,
        selected=selected,
        appendix_on=appendix_on,
        artist=artist_q,
        tracks=tracks,
        placeholders=placeholders
    )

@app.route("/all-in-one", methods=["GET"])
def all_in_one():
    # 1) список шаблонов БЕЗ файла приложения
    all_tmpls = list_templates()
    templates = [t for t in all_tmpls if t != APPENDIX_TEMPLATE_NAME]

    # 2) используем request.values (combo GET+POST) -> позволяет сохранять введённые значения
    vals = request.values

    selected = vals.getlist("t")
    selected = [s for s in selected if s in templates]

    appendix_on = vals.get("appendix") == "on"
    artist_q = (vals.get("artist") or "").strip()

    # 3) собираем плейсхолдеры: выбранные договоры + (опц.) приложение
    ph_paths = [os.path.join(TEMPLATE_DIR, f) for f in selected]
    if appendix_on:
        appx_path = os.path.join(TEMPLATE_DIR, APPENDIX_TEMPLATE_NAME)
        if os.path.exists(appx_path):
            ph_paths.append(appx_path)
    placeholders = extract_placeholders(ph_paths) if ph_paths else []

    # 4) если приложение включено, подтянем треки
    tracks = None
    if appendix_on and artist_q:
        try:
            tracks = rows_for_artist(artist_q).reset_index(drop=True)
        except Exception as e:
            flash(f"Ошибка чтения Excel: {e}", "warning")

    # 5) передаём в шаблон текущие значения (для предзаполнения)
    return render_template(
        "all_in_one.html",
        templates=templates,
        selected=selected,
        appendix_on=appendix_on,
        artist=artist_q,
        tracks=tracks,
        placeholders=placeholders,
        values=vals  # <— тут все текущие введённые поля, включая ph:...
    )


    # 1) Генерация договоров (каждый выбранный шаблон)
    for name in selected:
        path = os.path.join(TEMPLATE_DIR, name)
        doc = Document(path)
        replace_in_doc(doc, mapping)

        base = os.path.splitext(name)[0]
        base_lower = base.lower()
        if "договор" in base_lower:
            kind = "договор"
        elif "приложение" in base_lower:
            kind = "приложение"
        else:
            parts = re.split(r"[\s_()-]+", base)
            kind = (parts[-1] if parts and parts[-1] else base)

        out_name = f"{safe_fio} {kind}.docx"
        out_path = os.path.join(GENERATED_DIR, out_name)
        doc.save(out_path)
        out_files.append(out_path)

    # 2) Генерация «Приложение 1» (если включено)
    if appendix_on:
        appx_path = os.path.join(TEMPLATE_DIR, APPENDIX_TEMPLATE_NAME)
        if not os.path.exists(appx_path):
            flash(f"Не найден шаблон приложения: {APPENDIX_TEMPLATE_NAME}", "warning")
            return redirect(url_for("all_in_one"))

        # какие треки выбраны
        rows = rows_for_artist(artist_q).reset_index(drop=True) if artist_q else load_catalog().iloc[0:0]
        sel_idx = request.form.getlist("sel")
        try:
            idxs = [int(i) for i in sel_idx]
            idxs = [i for i in idxs if 0 <= i < len(rows)]
        except Exception:
            idxs = []
        if idxs:
            rows = rows.iloc[idxs]

        doc = Document(appx_path)
        # сначала плейсхолдеры по тексту приложения
        replace_in_doc(doc, mapping)
        # затем таблица с треками (нумерация, все колонки)
        fill_appendix_table(doc, rows)

        out_name = f"{safe_fio} приложение.docx"
        out_path = os.path.join(GENERATED_DIR, out_name)
        doc.save(out_path)
        out_files.append(out_path)

    # Отдаём один файл напрямую, или ZIP если файлов несколько
    if len(out_files) == 1:
        fp = out_files[0]
        return send_file(fp, as_attachment=True, download_name=os.path.basename(fp))

    mem = io.BytesIO()
    with ZipFile(mem, "w") as z:
        for fp in out_files:
            z.write(fp, arcname=os.path.basename(fp))
    mem.seek(0)
    zip_name = f"{safe_fio}.zip"
    return send_file(mem, as_attachment=True, download_name=zip_name)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
